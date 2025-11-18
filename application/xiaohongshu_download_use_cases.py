"""小红书图片下载用例实现"""

import asyncio
import re
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO
from loguru import logger
from PIL import Image

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from domain.exceptions import DomainError
from infrastructure.adapters.xiaohongshu_service import XiaohongshuBrowserService
from infrastructure.utils.word_generator import WordImageGenerator
from application.dto import XiaohongshuDownloadImagesCommand


class XiaohongshuDownloadImagesUseCase:
    """小红书图片下载用例"""

    def __init__(self, browser_service: XiaohongshuBrowserService):
        """初始化用例

        Args:
            browser_service: 小红书浏览器服务
        """
        self._browser_service = browser_service

    def _load_urls(self, command: XiaohongshuDownloadImagesCommand) -> List[str]:
        """加载 URL 列表

        Args:
            command: 下载命令

        Returns:
            URL 列表
        """
        urls = []

        # 从直接指定的 URL 列表加载
        if command.urls:
            urls.extend(command.urls)

        # 从配置文件加载
        if command.url_config_file:
            config_path = Path(command.url_config_file)
            if not config_path.exists():
                raise DomainError(f"URL config file not found: {config_path}")

            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith("#"):  # 忽略空行和注释
                            urls.append(line)
            except Exception as e:
                raise DomainError(f"Failed to load URL config file: {e}")

        if not urls:
            raise DomainError("No URLs provided. Use --urls or --url-config-file")

        # 去重
        unique_urls = list(dict.fromkeys(urls))  # 保持顺序的去重
        logger.info(f"Loaded {len(unique_urls)} unique URLs")
        return unique_urls

    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名，移除不合法的字符

        Args:
            filename: 原始文件名

        Returns:
            清理后的文件名
        """
        # 移除或替换 Windows/Linux 文件系统不支持的字符
        # 不允许的字符: < > : " / \ | ? *
        illegal_chars = r'[<>:"/\\|?*]'
        sanitized = re.sub(illegal_chars, "_", filename)
        # 移除前后空格和点
        sanitized = sanitized.strip(" .")
        return sanitized

    def _generate_filename(self, title: str) -> str:
        """生成 Word 文件名：日期 + title 前20个字符

        Args:
            title: 页面标题

        Returns:
            生成的文件名（不含扩展名）
        """
        # 获取当前日期，格式：YYYYMMDD
        date_str = datetime.now().strftime("%Y%m%d")

        # 提取 title 前20个字符
        title_prefix = title[:20] if title else "小红书"
        # 清理不合法的文件名字符
        title_prefix = self._sanitize_filename(title_prefix)

        # 如果清理后为空，使用默认值
        if not title_prefix:
            title_prefix = "小红书"

        # 组合文件名
        filename = f"{date_str}{title_prefix}"
        return filename

    async def _download_images_for_url(self, url: str) -> List[Dict[str, Any]]:
        """为单个 URL 下载图片

        Args:
            url: 内容页面 URL

        Returns:
            图片信息列表，每个元素包含 {'url': str, 'index': int, 'data': bytes}
        """
        logger.info(f"Processing URL: {url}")

        # 获取图片 URL 列表
        images_info = await self._browser_service.get_content_images(url)

        if not images_info:
            logger.warning(f"No images found for URL: {url}")
            return []

        # 下载图片
        word_generator = WordImageGenerator()
        images_with_data = []

        for img_info in images_info:
            img_url = img_info["url"]
            img_index = img_info["index"]

            try:
                # 下载图片
                image_data = await word_generator.download_image(img_url)

                # 确保图片格式是 Word 支持的格式
                image_data = self._ensure_word_compatible_format(image_data)

                images_with_data.append(
                    {
                        "url": img_url,
                        "index": img_index,
                        "data": image_data,
                    }
                )
                logger.info(f"Downloaded and processed image {img_index} from {url}")
            except Exception as e:
                logger.error(f"Failed to download or process image {img_index} from {img_url}: {e}")
                continue

        return images_with_data

    async def execute(self, command: XiaohongshuDownloadImagesCommand) -> str:
        """执行图片下载用例

        Args:
            command: 下载命令

        Returns:
            生成的 Word 文件路径
        """
        logger.info("Starting Xiaohongshu download images use case")

        # 加载 URL 列表
        urls = self._load_urls(command)

        # 获取第一个 URL 的 title 用于生成文件名
        page_title = ""
        if urls:
            try:
                page_title = await self._browser_service.get_page_title(urls[0])
                logger.info(f"Got page title from first URL: {page_title}")
            except Exception as e:
                logger.warning(f"Failed to get page title from first URL: {e}")

        # 生成文件名
        filename_base = self._generate_filename(page_title)
        filename = f"{filename_base}.docx"
        logger.info(f"Generated filename: {filename}")

        # 创建 Word 文档
        doc = Document()

        # 设置页面边距为 0
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

        # 逐个处理 URL
        for url in urls:
            try:
                # 获取并下载图片
                images_with_data = await self._download_images_for_url(url)

                if not images_with_data:
                    logger.warning(f"No images downloaded for URL: {url}")
                    continue

                # 按 index 排序（虽然已经排序，但确保一下）
                images_with_data.sort(key=lambda x: x["index"])

                # 添加图片到文档
                for img_info in images_with_data:
                    try:
                        image_data = img_info["data"]
                        if not image_data:
                            logger.warning(f"Image {img_info['index']} data is empty, skipping")
                            continue
                        self._add_image_to_document(doc, image_data)
                        logger.debug(f"Added image {img_info['index']} to document")
                    except Exception as e:
                        logger.error(
                            f"Error adding image {img_info['index']} to document: {type(e).__name__}: {str(e)}"
                        )
                        logger.exception("Full traceback:")
                        continue

                # 在处理下一个 URL 前稍作等待
                await asyncio.sleep(1)

            except Exception as e:
                logger.error(f"Error processing URL {url}: {e}")
                continue

        # 保存文档
        output_path = Path(command.output_dir) / filename
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        logger.info(f"Word document saved to: {output_path}")
        return str(output_path)

    def _detect_image_format(self, image_data: bytes) -> str:
        """通过文件头检测图片真实格式

        Args:
            image_data: 图片二进制数据

        Returns:
            图片格式字符串（如 'JPEG', 'PNG', 'WEBP', 'GIF', 'BMP'）
        """
        if len(image_data) < 12:
            return "UNKNOWN"

        # 检查文件头（magic bytes）
        header = image_data[:12]

        # JPEG: FF D8 FF
        if header[:3] == b"\xff\xd8\xff":
            return "JPEG"

        # PNG: 89 50 4E 47 0D 0A 1A 0A
        if header[:8] == b"\x89\x50\x4e\x47\x0d\x0a\x1a\x0a":
            return "PNG"

        # GIF: 47 49 46 38 (GIF8)
        if header[:4] == b"\x47\x49\x46\x38":
            return "GIF"

        # BMP: 42 4D (BM)
        if header[:2] == b"\x42\x4d":
            return "BMP"

        # WebP: RIFF...WEBP
        if header[:4] == b"RIFF" and len(image_data) >= 12:
            if image_data[8:12] == b"WEBP":
                return "WEBP"

        # AVIF: ftyp...avif
        if header[4:8] in [b"avif", b"avis"]:
            return "AVIF"

        return "UNKNOWN"

    def _convert_to_png(self, image_data: bytes) -> bytes:
        """将图片转换为 PNG 格式

        Args:
            image_data: 原始图片二进制数据

        Returns:
            PNG 格式的图片二进制数据
        """
        img = Image.open(BytesIO(image_data))

        # 如果图片有透明通道（RGBA），保持透明
        # 否则转换为 RGB
        if img.mode in ("RGBA", "LA", "P"):
            # 保持透明通道
            output = BytesIO()
            img.save(output, format="PNG")
        else:
            # 转换为 RGB 再保存为 PNG
            rgb_img = img.convert("RGB")
            output = BytesIO()
            rgb_img.save(output, format="PNG")

        return output.getvalue()

    def _ensure_word_compatible_format(self, image_data: bytes) -> bytes:
        """确保图片格式是 Word 支持的格式，如果不支持则转换为 PNG

        Word 支持的格式：JPEG, PNG, GIF, BMP
        不支持的格式：WebP, AVIF 等

        Args:
            image_data: 原始图片二进制数据

        Returns:
            Word 兼容格式的图片二进制数据
        """
        format_name = self._detect_image_format(image_data)

        # Word 支持的格式
        word_supported_formats = {"JPEG", "PNG", "GIF", "BMP"}

        if format_name in word_supported_formats:
            logger.debug(f"Image format {format_name} is Word-compatible, keeping original")
            return image_data

        # 不支持的格式，转换为 PNG
        logger.info(f"Image format {format_name} is not Word-compatible, converting to PNG")
        try:
            png_data = self._convert_to_png(image_data)
            logger.info(f"Successfully converted {format_name} to PNG")
            return png_data
        except Exception as e:
            logger.error(f"Failed to convert {format_name} to PNG: {e}")
            # 如果转换失败，尝试用 PIL 直接打开并保存
            try:
                img = Image.open(BytesIO(image_data))
                output = BytesIO()
                img.save(output, format="PNG")
                return output.getvalue()
            except Exception as e2:
                logger.error(f"Failed to convert image using PIL: {e2}")
                raise ValueError(f"Cannot convert {format_name} image to PNG") from e2

    def _get_image_size(self, image_data: bytes) -> tuple[int, int]:
        """获取图片尺寸

        Args:
            image_data: 图片二进制数据

        Returns:
            (width, height) 元组
        """
        img = Image.open(BytesIO(image_data))
        return img.size

    def _get_page_dimensions(self, doc: Document) -> tuple[float, float]:
        """获取 Word 页面的可用尺寸（英寸）

        Args:
            doc: Word 文档对象

        Returns:
            (width, height) 元组，单位为英寸
        """
        section = doc.sections[0]
        # 页面宽度 = 页面总宽度 - 左边距 - 右边距
        page_width = (
            section.page_width.inches - section.left_margin.inches - section.right_margin.inches
        )
        # 页面高度 = 页面总高度 - 上边距 - 下边距
        page_height = (
            section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
        )

        return page_width, page_height

    def _add_image_to_document(
        self, doc: Document, image_data: bytes, max_width: float = 6.5
    ) -> None:
        """向文档添加图片（居中），图片等比缩放以适配页面大小

        图片会等比缩放，使得高度或宽度至少一个和页面大小一样（填满页面）

        Args:
            doc: Word 文档对象
            image_data: 图片二进制数据
            max_width: 已废弃，保留用于兼容性

        Raises:
            Exception: 如果图片处理或添加失败
        """
        # 验证图片数据
        if not image_data:
            raise ValueError("Image data is empty")

        # 获取图片尺寸（像素）
        try:
            width_px, height_px = self._get_image_size(image_data)
        except Exception as e:
            raise ValueError(f"Failed to get image size: {e}") from e

        # 验证尺寸有效性
        if width_px <= 0 or height_px <= 0:
            raise ValueError(f"Invalid image dimensions: {width_px}x{height_px}")

        # 假设图片 DPI 为 96，将像素转换为英寸
        # 如果图片有 EXIF 信息，可以使用实际 DPI
        dpi = 96.0
        width_inches = width_px / dpi
        height_inches = height_px / dpi

        # 获取页面可用尺寸
        page_width, page_height = self._get_page_dimensions(doc)
        logger.debug(f"Page dimensions: {page_width:.2f}x{page_height:.2f} inches")
        logger.debug(f"Image dimensions: {width_inches:.2f}x{height_inches:.2f} inches")

        # 计算两个缩放比例（按宽度和按高度）
        scale_by_width = page_width / width_inches
        scale_by_height = page_height / height_inches

        # 选择较大的缩放比例，这样至少一个维度会填满页面
        # 同时保持图片的宽高比
        scale = max(scale_by_width, scale_by_height)

        # 应用缩放比例
        scaled_width = width_inches * scale
        scaled_height = height_inches * scale

        # 确保缩放后的尺寸不超过页面大小（理论上不应该超过，但做安全检查）
        if scaled_width > page_width:
            # 如果宽度超过，按宽度重新缩放
            scale = page_width / width_inches
            scaled_width = page_width
            scaled_height = height_inches * scale
            logger.debug(
                f"Width exceeded page, rescaling by width: {scaled_width:.2f}x{scaled_height:.2f}"
            )
        elif scaled_height > page_height:
            # 如果高度超过，按高度重新缩放
            scale = page_height / height_inches
            scaled_width = width_inches * scale
            scaled_height = page_height
            logger.debug(
                f"Height exceeded page, rescaling by height: {scaled_width:.2f}x{scaled_height:.2f}"
            )

        logger.debug(
            f"Final scaled image dimensions: {scaled_width:.2f}x{scaled_height:.2f} inches"
        )

        # 验证缩放后的尺寸
        if scaled_width <= 0 or scaled_height <= 0:
            raise ValueError(f"Invalid scaled dimensions: {scaled_width}x{scaled_height}")

        # 添加段落并居中
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图片（使用宽度，高度会自动按比例调整）
        try:
            run = para.add_run()
            run.add_picture(BytesIO(image_data), width=Inches(scaled_width))
        except Exception as e:
            raise RuntimeError(f"Failed to add picture to document: {e}") from e

        # 不添加空行，避免产生空白页
        # 图片已经填满页面，不需要额外的空行
