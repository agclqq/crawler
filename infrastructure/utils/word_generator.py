"""Word 文档生成工具"""

import os
import httpx
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO
from loguru import logger

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


class WordImageGenerator:
    """Word 图片文档生成器"""

    def __init__(self, output_dir: str = "output"):
        """初始化生成器

        Args:
            output_dir: 输出目录
        """
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def download_image(self, url: str, timeout: int = 30) -> bytes:
        """下载图片

        Args:
            url: 图片 URL
            timeout: 超时时间（秒）

        Returns:
            图片的二进制数据
        """
        logger.info(f"Downloading image: {url}")
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.get(url, follow_redirects=True)
            response.raise_for_status()
            return response.content

    def _get_image_size(self, image_data: bytes) -> tuple[int, int]:
        """获取图片尺寸

        Args:
            image_data: 图片二进制数据

        Returns:
            (width, height) 元组
        """
        img = Image.open(BytesIO(image_data))
        return img.size

    def create_word_document(
        self, url_images_map: Dict[str, List[Dict[str, Any]]], filename: str = "xiaohongshu_images.docx"
    ) -> str:
        """创建包含图片的 Word 文档

        Args:
            url_images_map: URL 到图片列表的映射，格式为 {url: [{'url': str, 'index': int}, ...]}
            filename: 输出文件名

        Returns:
            生成的 Word 文件路径
        """
        doc = Document()

        # 设置页面边距为 0
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0)
            section.right_margin = Inches(0)

        # 遍历每个 URL 的图片
        for url, images in url_images_map.items():
            if not images:
                logger.warning(f"No images found for URL: {url}")
                continue

            logger.info(f"Processing {len(images)} images from URL: {url}")

            # 添加 URL 作为标题（可选，如果需要的话可以注释掉）
            # para = doc.add_paragraph(url)
            # para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # run = para.runs[0] if para.runs else para.add_run()
            # run.font.size = Pt(10)

            # 添加图片
            for img_info in images:
                img_url = img_info["url"]
                img_index = img_info.get("index", 0)

                try:
                    # 下载图片（这里需要同步调用，因为 docx 库是同步的）
                    # 注意：在实际使用中，图片应该在用例中先下载好
                    # 这里假设图片数据已经准备好
                    logger.debug(f"Adding image {img_index} from {img_url}")

                    # 这里暂时跳过，因为需要在异步上下文中下载
                    # 实际实现中，图片应该在用例中下载并传递给这个方法

                except Exception as e:
                    logger.error(f"Error adding image {img_index}: {e}")
                    continue

        output_path = self.output_dir / filename
        doc.save(str(output_path))
        logger.info(f"Word document saved to: {output_path}")
        return str(output_path)

    def add_image_to_document(self, doc: Document, image_data: bytes, max_width: float = 6.5) -> None:
        """向文档添加图片（居中）

        Args:
            doc: Word 文档对象
            image_data: 图片二进制数据
            max_width: 最大宽度（英寸）
        """
        # 获取图片尺寸
        width, height = self._get_image_size(image_data)

        # 计算缩放比例，使宽度不超过 max_width
        scale = max_width / (width / 96)  # 96 DPI 假设
        scaled_width = (width / 96) * scale
        scaled_height = (height / 96) * scale

        # 添加段落并居中
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加图片
        run = para.add_run()
        run.add_picture(BytesIO(image_data), width=Inches(scaled_width))

        # 添加空行
        doc.add_paragraph()

